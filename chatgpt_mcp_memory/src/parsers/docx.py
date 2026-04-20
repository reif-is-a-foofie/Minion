"""DOCX parser via python-docx."""
from __future__ import annotations

from pathlib import Path

from . import ParsedChunk, ParseResult
from ._common import chunk_text


def parse(path: Path) -> ParseResult:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n\n".join(parts)
    chunks = [
        ParsedChunk(text=t, role=None, meta={"seq": i})
        for i, t in enumerate(chunk_text(text))
    ]
    return ParseResult(
        chunks=chunks,
        source_meta={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
        kind="docx",
        parser="python-docx",
    )
